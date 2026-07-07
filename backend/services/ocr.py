import base64
import json
import os
import requests
import tempfile
import urllib3
from pathlib import Path
import ftfy

# Cloudflare/ngrok tunnels use certs Python's bundled CA bundle may not verify.
# Since these are trusted dev tunnels (user-controlled), we skip SSL verification.
urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)

def clean_text(text: str) -> str:
    """Fix mojibake (â€" → —, â€˜ → ', etc.) and normalise whitespace."""
    return ftfy.fix_text(text)

# ── EasyOCR ───────────────────────────────────────────────────────────────────
_easyocr_reader = None

def _get_easyocr():
    global _easyocr_reader
    if _easyocr_reader is None:
        import easyocr
        _easyocr_reader = easyocr.Reader(["en"], gpu=False)
    return _easyocr_reader

def _layout_text(results: list) -> str:
    """
    Reconstruct reading order from EasyOCR bounding boxes.
    Groups detections into visual lines by Y position, then sorts each
    line left-to-right by X — preserves section headers, bullets, and
    multi-column sidebars instead of flattening everything into one string.
    """
    if not results:
        return ""
    blocks = []
    for bbox, text, _conf in results:
        xs = [p[0] for p in bbox]
        ys = [p[1] for p in bbox]
        blocks.append({
            "y": (min(ys) + max(ys)) / 2,
            "x": min(xs),
            "h": max(ys) - min(ys),
            "t": text,
        })
    blocks.sort(key=lambda b: b["y"])
    median_h = sorted(b["h"] for b in blocks)[len(blocks) // 2]
    thresh = median_h * 0.6

    lines, cur = [], [blocks[0]]
    for b in blocks[1:]:
        if abs(b["y"] - cur[-1]["y"]) <= thresh:
            cur.append(b)
        else:
            lines.append(cur)
            cur = [b]
    lines.append(cur)

    return "\n".join(
        "  ".join(b["t"] for b in sorted(ln, key=lambda b: b["x"]))
        for ln in lines
    )

def extract_easyocr(file_bytes: bytes, filename: str) -> str:
    suffix = Path(filename).suffix.lower()
    with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as tmp:
        tmp.write(file_bytes)
        tmp_path = tmp.name
    try:
        reader = _get_easyocr()
        results = reader.readtext(tmp_path)
        return clean_text(_layout_text(results))
    finally:
        os.unlink(tmp_path)


def _ocr_image_bytes(img_bytes: bytes) -> str:
    """Run EasyOCR on raw PNG/JPEG bytes (no filename needed)."""
    with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp:
        tmp.write(img_bytes)
        tmp_path = tmp.name
    try:
        results = _get_easyocr().readtext(tmp_path)
        return _layout_text(results)
    finally:
        os.unlink(tmp_path)


def extract_pdf(file_bytes: bytes) -> str:
    """
    Extract text from a PDF.
    Tries embedded text first (fast, works for digital PDFs).
    Falls back to EasyOCR per page if the PDF is a scan / has no embedded text.
    """
    import fitz  # pymupdf
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    pages_text = [page.get_text() for page in doc]
    embedded = clean_text("\n".join(pages_text))
    if len(embedded.strip()) > 100:
        return embedded
    # Scanned PDF — render each page and OCR
    all_text = []
    doc2 = fitz.open(stream=file_bytes, filetype="pdf")  # re-open for rendering
    for page in doc2:
        pix = page.get_pixmap(dpi=200)
        all_text.append(_ocr_image_bytes(pix.tobytes("png")))
    return clean_text("\n".join(all_text))


def extract_docx(file_bytes: bytes) -> str:
    """Extract text from a DOCX file including table cells."""
    import io
    from docx import Document
    doc = Document(io.BytesIO(file_bytes))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text.strip())
    return clean_text("\n".join(parts))


def convert_docx_to_pdf(docx_bytes: bytes) -> bytes | None:
    """
    Convert DOCX bytes to PDF bytes using docx2pdf (requires Microsoft Word on Windows).
    Returns None if conversion is unavailable or fails.
    """
    import tempfile
    try:
        from docx2pdf import convert
    except ImportError:
        return None
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp.write(docx_bytes)
        tmp_path = tmp.name
    pdf_path = tmp_path[:-5] + ".pdf"
    try:
        convert(tmp_path, pdf_path)
        with open(pdf_path, "rb") as f:
            return f.read()
    except Exception:
        return None
    finally:
        try: os.unlink(tmp_path)
        except OSError: pass
        try: os.unlink(pdf_path)
        except OSError: pass


# ── MiniCPM-V via Ollama ──────────────────────────────────────────────────────
OCR_SYSTEM_PROMPT = (
    "You are an OCR engine. Transcribe ALL visible text from this resume image exactly as it "
    "appears, reading in natural order (top to bottom, left to right within columns). Output "
    "only the raw extracted text — no commentary, no markdown, nothing that isn't visibly in the image."
)

def _ollama_chat(base_url: str, model: str, image_b64: str, timeout: int = 300) -> str:
    payload = {
        "model": model,
        "messages": [
            {"role": "system", "content": OCR_SYSTEM_PROMPT},
            {"role": "user", "content": "Extract all text from this resume image.", "images": [image_b64]},
        ],
        "stream": True,
        "options": {"num_ctx": 16384, "num_predict": -1},
    }
    response = requests.post(f"{base_url.rstrip('/')}/api/chat", json=payload, stream=True, timeout=timeout, verify=False)
    response.raise_for_status()
    content = ""
    for line in response.iter_lines(decode_unicode=True):
        if line:
            chunk = json.loads(line)
            content += chunk.get("message", {}).get("content", "")
    return clean_text(content.strip())

def _resize_image(file_bytes: bytes, max_px: int = 1600) -> bytes:
    from PIL import Image
    import io
    img = Image.open(io.BytesIO(file_bytes)).convert("RGB")
    w, h = img.size
    if max(w, h) > max_px:
        scale = max_px / max(w, h)
        img = img.resize((int(w * scale), int(h * scale)), Image.LANCZOS)
    buf = io.BytesIO()
    img.save(buf, format="JPEG", quality=85)
    return buf.getvalue()

def _poll_job(base_url: str, job_id: str, interval: int = 5, max_wait: int = 600) -> str:
    """Poll /result/{job_id} until done or error. Tolerates transient tunnel blips."""
    import time
    deadline = time.time() + max_wait
    while time.time() < deadline:
        try:
            r = requests.get(f"{base_url.rstrip('/')}/result/{job_id}", verify=False, timeout=10)
            r.raise_for_status()
            data = r.json()
        except (requests.exceptions.RequestException, ValueError):
            time.sleep(interval)
            continue
        if data["status"] == "done":
            return data["result"]
        if data["status"] == "error":
            raise RuntimeError(f"Inference failed on Colab: {data['result']}")
        time.sleep(interval)
    raise TimeoutError(f"Inference did not complete within {max_wait}s")

def extract_minicpm(file_bytes: bytes, ollama_url: str, model: str = None) -> str:
    model = model or os.getenv("MINICPM_MODEL", "minicpm-v")
    compressed = _resize_image(file_bytes)
    image_b64 = base64.b64encode(compressed).decode("utf-8")

    base_url = ollama_url.rstrip("/")
    # Try async job queue first (Colab FastAPI server)
    try:
        submit_res = requests.post(
            f"{base_url}/submit",
            json={"image_b64": image_b64, "model": model},
            verify=False, timeout=15,
        )
        if submit_res.status_code == 200:
            job_id = submit_res.json()["job_id"]
            return clean_text(_poll_job(base_url, job_id))
    except (requests.exceptions.ConnectionError, KeyError):
        pass  # fall back to direct Ollama streaming

    # Fallback: direct Ollama /api/chat (works with ngrok, times out with Cloudflare)
    return _ollama_chat(ollama_url, model, image_b64)
